#!/usr/bin/env python3
"""Build a BMC-formatted .docx from the manuscript markdown.
Figures are submitted separately per BMC policy; legends remain in the text.
"""
import re, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = os.path.join(os.path.dirname(__file__), "..", "manuscript", "TR_ODCL_manuscript.md")
OUT = os.path.join(os.path.dirname(__file__), "..", "manuscript", "TR_ODCL_manuscript.docx")

doc = Document()
# Base style: Times New Roman 12, double-ish spacing acceptable
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.5

NAVY = RGBColor(0x0b, 0x3d, 0x66)

def add_runs(p, text):
    """Split on ** for bold; leave rest plain."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for seg in parts:
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = p.add_run(seg[2:-2]); r.bold = True
        else:
            p.add_run(seg)

lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
table_buffer = []
in_table = False

def flush_table():
    global table_buffer
    if not table_buffer:
        return
    # parse markdown table rows
    rows = [r for r in table_buffer if r.strip().startswith("|")]
    # drop separator row (---)
    parsed = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        if all(set(c) <= set("-: ") for c in cells):
            continue
        parsed.append(cells)
    if parsed:
        ncol = len(parsed[0])
        t = doc.add_table(rows=len(parsed), cols=ncol)
        t.style = "Table Grid"  # BMC: no colour/shading permitted in tables
        for ri, row in enumerate(parsed):
            for ci, val in enumerate(row):
                cell = t.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs(p, val)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    if ri == 0:
                        run.bold = True
    table_buffer = []

for raw in lines:
    line = raw.rstrip("\n")
    stripped = line.strip()
    # table accumulation
    if stripped.startswith("|"):
        in_table = True
        table_buffer.append(line)
        continue
    elif in_table:
        flush_table()
        in_table = False
    if stripped == "" or stripped == "---":
        continue
    if stripped.startswith("# "):
        p = doc.add_heading(level=0)
        r = p.add_run(stripped[2:]); r.font.color.rgb = NAVY
        continue
    if stripped.startswith("### "):
        p = doc.add_heading(level=2)
        r = p.add_run(stripped[4:]); r.font.color.rgb = NAVY
        continue
    if stripped.startswith("## "):
        p = doc.add_heading(level=1)
        r = p.add_run(stripped[3:]); r.font.color.rgb = NAVY
        continue
    if stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, stripped[2:])
        continue
    # numbered reference lines "1. ..."
    m = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if m:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        add_runs(p, f"{m.group(1)}. {m.group(2)}")
        continue
    p = doc.add_paragraph()
    add_runs(p, stripped)

if in_table:
    flush_table()

doc.save(OUT)
print("built", OUT, os.path.getsize(OUT), "bytes")
